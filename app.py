from flask import Flask, render_template, request, jsonify, send_file
import os
import sys
import json
import threading
import time
import logging
import io
from datetime import datetime

# Add imports for ZIP file creation
import zipfile
from io import BytesIO

# Import GoogleTranslator for text translation
from deep_translator import GoogleTranslator

# Import Hugging Face pipeline for sentiment analysis
from transformers import pipeline

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__, static_folder='.', static_url_path='')

# Initialize the sentiment analysis pipeline
tone_analyzer = pipeline("sentiment-analysis")

# Import the speech recognition module
try:
    from speech_translator import SpeechTranslator
    translator = SpeechTranslator()
    speech_module_available = True
    logger.info("Speech translator module loaded successfully")
except ImportError as e:
    logger.error(f"Error importing speech translator: {e}")
    speech_module_available = False
    translator = None

# Try to import document generation libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    reportlab_available = True
    logger.info("ReportLab (PDF generation) loaded successfully")
except ImportError as e:
    logger.error(f"Error importing ReportLab: {e}")
    reportlab_available = False

try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    docx_available = True
    logger.info("python-docx (DOCX generation) loaded successfully")
except ImportError as e:
    logger.error(f"Error importing python-docx: {e}")
    docx_available = False

# Store recognition results
recognition_results = {}

# Define the upload folder
UPLOAD_FOLDER = 'uploads'

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/dashboard')
def dashboard():
    return app.send_static_file('dashboard.html')

# Add a route to handle audio file uploads
@app.route('/api/upload-audio', methods=['POST'])
def upload_audio():
    audio_file = request.files.get('audio')
    language = request.form.get('language', 'en')

    if not audio_file:
        return jsonify({'success': False, 'error': 'No audio file provided'}), 400

    # Validate file type
    if not audio_file.filename.lower().endswith(('.wav', '.mp3')):
        return jsonify({'success': False, 'error': 'Unsupported file format. Please upload a .wav or .mp3 file.'}), 400

    # Save the uploaded file
    file_path = os.path.join(UPLOAD_FOLDER, audio_file.filename)
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    audio_file.save(file_path)

    try:
        # Transcribe the audio
        original_text = translator.transcribe_audio_file(file_path, language=language)
        if not original_text:
            return jsonify({'success': False, 'error': 'Could not transcribe audio'}), 500

        return jsonify({'success': True, 'transcription': original_text})
    except Exception as e:
        logger.error(f"Error uploading audio: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/recognize', methods=['POST'])
def recognize_speech():
    if not speech_module_available:
        return jsonify({
            'success': False,
            'error': 'Speech recognition module not available'
        }), 500

    try:
        # Recognize speech
        text = translator.recognize_speech(language="en-US")
        if text:
            return jsonify({
                'success': True,
                'transcription': text
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Could not recognize speech'
            }), 500
    except Exception as e:
        logger.error(f"Error recognizing speech: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Update the get_results route to include tone and summary
@app.route('/api/results/<session_id>', methods=['GET'])
def get_results(session_id):
    if session_id in recognition_results:
        result = recognition_results[session_id]
        
        if 'error' in result:
            return jsonify({
                'success': False,
                'error': result['error']
            })
        
        return jsonify({
            'success': True,
            'original': result['original'],
            'translated': result['translated'],
            'src_lang': result['src_lang'],
            'dest_lang': result['dest_lang'],
            'src_lang_name': result.get('src_lang_name', 'Unknown'),
            'dest_lang_name': result.get('dest_lang_name', 'Unknown'),
            'timestamp': result.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            'tone': result.get('tone', 'Neutral'),
            'summary': result.get('summary', '')
        })
    
    return jsonify({
        'success': False,
        'error': 'Results not found or recognition still in progress'
    })

@app.route('/api/speak', methods=['POST'])
def speak_text():
    if not speech_module_available:
        return jsonify({
            'success': False,
            'error': 'Speech module not available'
        }), 500

    data = request.json
    text = data.get('text', '')
    language = data.get('language', 'en')

    if not text:
        return jsonify({
            'success': False,
            'error': 'No text provided'
        }), 400

    try:
        # Use the SpeechTranslator module to speak the text
        translator.speak_text(text, language=language)

        return jsonify({
            'success': True,
            'message': 'Text spoken successfully'
        })
    except Exception as e:
        logger.error(f"Error speaking text: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/languages', methods=['GET'])
def get_languages():
    if not speech_module_available:
        return jsonify({
            'success': False,
            'error': 'Speech module not available'
        }), 500
    
    return jsonify({
        'success': True,
        'languages': translator.languages
    })

@app.route('/api/generate-pdf', methods=['POST'])
def generate_pdf():
    if not reportlab_available:
        return jsonify({
            'success': False,
            'error': 'PDF generation module not available'
        }), 500
    
    data = request.json
    text = data.get('text', '')
    filename = data.get('filename', 'document')
    
    if not text:
        return jsonify({
            'success': False,
            'error': 'No text provided'
        }), 400
    
    try:
        # Create a PDF in memory
        buffer = io.BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        normal_style = styles['Normal']
        normal_style.fontSize = 10
        normal_style.leading = 14
        
        # Create the content
        content = []
        
        # Add title
        content.append(Paragraph(f"Speech Recognition Results", title_style))
        content.append(Spacer(1, 12))
        
        # Add timestamp
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        content.append(Paragraph(f"Generated on: {timestamp}", styles['Italic']))
        content.append(Spacer(1, 12))
        
        # Add the text content
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                content.append(Paragraph(para, normal_style))
                content.append(Spacer(1, 6))
        
        # Build the PDF
        doc.build(content)
        
        # Get the value from the buffer
        pdf_value = buffer.getvalue()
        buffer.close()
        
        # Create a response
        response_buffer = io.BytesIO(pdf_value)
        
        return send_file(
            response_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"{filename}.pdf"
        )
        
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/generate-docx', methods=['POST'])
def generate_docx():
    if not docx_available:
        return jsonify({
            'success': False,
            'error': 'DOCX generation module not available'
        }), 500
    
    data = request.json
    text = data.get('text', '')
    filename = data.get('filename', 'document')
    
    if not text:
        return jsonify({
            'success': False,
            'error': 'No text provided'
        }), 400
    
    try:
        # Create a DOCX document
        document = docx.Document()
        
        # Add title
        title = document.add_heading('Speech Recognition Results', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        timestamp_para = document.add_paragraph(f"Generated on: {timestamp}")
        timestamp_para.style = 'Italic'
        
        # Add the text content
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                document.add_paragraph(para)
        
        # Save the document to a BytesIO object
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{filename}.docx"
        )
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Add a route to download both PDF and DOCX as a ZIP file
@app.route('/api/download-all', methods=['POST'])
def download_all():
    data = request.json
    text = data.get('text', '')
    filename = data.get('filename', 'document')
    audio_file_path = data.get('audio_file_path', '')  # Path to the uploaded audio file

    if not text:
        return jsonify({
            'success': False,
            'error': 'No text provided'
        }), 400

    try:
        # Create a ZIP file in memory
        memory_file = BytesIO()
        with zipfile.ZipFile(memory_file, 'w') as zf:
            # Add PDF file
            if reportlab_available:
                pdf_buffer = BytesIO()
                doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
                styles = getSampleStyleSheet()

                title_style = ParagraphStyle(
                    'Title',
                    parent=styles['Heading1'],
                    alignment=TA_CENTER,
                    spaceAfter=12
                )

                normal_style = styles['Normal']
                normal_style.fontSize = 10
                normal_style.leading = 14

                content = []
                content.append(Paragraph(f"Speech Recognition Results", title_style))
                content.append(Spacer(1, 12))

                timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                content.append(Paragraph(f"Generated on: {timestamp}", styles['Italic']))
                content.append(Spacer(1, 12))

                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        content.append(Paragraph(para, normal_style))
                        content.append(Spacer(1, 6))

                doc.build(content)
                pdf_value = pdf_buffer.getvalue()
                zf.writestr(f"{filename}.pdf", pdf_value)

            # Add DOCX file
            if docx_available:
                docx_buffer = BytesIO()
                document = docx.Document()

                title = document.add_heading('Speech Recognition Results', level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                timestamp_para = document.add_paragraph(f"Generated on: {timestamp}")
                timestamp_para.style = 'Italic'

                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        document.add_paragraph(para)

                document.save(docx_buffer)
                docx_value = docx_buffer.getvalue()
                zf.writestr(f"{filename}.docx", docx_value)

            # Add audio file
            if os.path.exists(audio_file_path):
                with open(audio_file_path, 'rb') as audio_file:
                    zf.writestr(os.path.basename(audio_file_path), audio_file.read())

        # Prepare the ZIP file for download
        memory_file.seek(0)

        return send_file(
            memory_file,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f"{filename}_files.zip"
        )

    except Exception as e:
        logger.error(f"Error creating ZIP file: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Add a new route to process audio directly
@app.route('/api/process-audio', methods=['POST'])
def process_audio():
    try:
        logger.info("Processing audio file...")
        if not speech_module_available:
            logger.error("Speech module not available")
            return jsonify({'success': False, 'error': 'Speech module not available'}), 500

        audio_file = request.files.get('audio')
        if not audio_file:
            logger.error("No audio file provided")
            return jsonify({'success': False, 'error': 'No audio file provided'}), 400

        file_path = os.path.join(UPLOAD_FOLDER, audio_file.filename)
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        audio_file.save(file_path)
        logger.info(f"Audio file saved to {file_path}")

        speech_lang = request.form.get('language', 'en-US')
        original_text = translator.transcribe_audio_file(file_path, language=speech_lang)
        if not original_text:
            logger.error("Could not transcribe audio")
            return jsonify({'success': False, 'error': 'Could not transcribe audio'}), 500

        logger.info("Audio transcription successful")
        return jsonify({'success': True, 'transcription': original_text})
    except Exception as e:
        logger.error(f"Error processing audio: {e}")
        # Add a proper return statement for the exception case
        return jsonify({'success': False, 'error': str(e)}), 500

# Add a route for summarization
@app.route('/api/summarize', methods=['POST'])
def summarize_text():
    data = request.json
    text = data.get('text', '')

    if not text.strip():
        return jsonify({'success': False, 'error': 'No text provided'}), 400

    try:
        if speech_module_available and hasattr(translator, 'generate_summary'):
            summary = translator.generate_summary(text)
            return jsonify({'success': True, 'summary': summary})
        else:
            # Fallback summarization logic
            sentences = text.split('.')
            if len(sentences) <= 3:
                summary = text
            else:
                # Summarization logic
                summary = '. '.join(sentences[:3]) + '.'
            return jsonify({'success': True, 'summary': summary})
    except Exception as e:
        logger.error(f"Error generating summary: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/translate', methods=['POST'])
def translate():
    data = request.json
    text = data.get('text', '')
    src_lang = data.get('src_lang', 'auto')
    dest_lang = data.get('dest_lang', 'en')

    if not text:
        return jsonify({'success': False, 'error': 'No text provided'}), 400

    try:
        translated_text = translator.translate_text(text, src=src_lang, dest=dest_lang)
        if not translated_text:
            return jsonify({'success': False, 'error': 'Translation failed'}), 500

        return jsonify({'success': True, 'translated_text': translated_text})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/analyze-tone', methods=['POST'])
def analyze_tone():
    try:
        data = request.json
        text = data.get('text', '')

        if not text:
            return jsonify({'success': False, 'error': 'No text provided'}), 400

        # Analyze the tone using the Hugging Face pipeline
        result = tone_analyzer(text)

        # Extract the label (e.g., POSITIVE, NEGATIVE, NEUTRAL)
        tone = result[0]['label']

        return jsonify({'success': True, 'tone': tone})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    # Check if the speech recognition module is available
    if not speech_module_available:
        print("WARNING: Speech recognition module not available. Some features will be disabled.")
    
    # Check if document generation modules are available
    if not reportlab_available:
        print("WARNING: PDF generation module not available. PDF export will be disabled.")
    
    if not docx_available:
        print("WARNING: DOCX generation module not available. DOCX export will be disabled.")
    
    # Get port from environment variable or use default
    port = int(os.environ.get('PORT', 5000))
    
    # Run the Flask app
    app.run(host='0.0.0.0', port=port, debug=True)